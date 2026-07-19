"""Aplica formato editorial reproducible al informe DOCX generado por Pandoc."""

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
BORDER = "D7DBE2"
USABLE_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, size, color="000000", bold=False, before=0, after=6, line=1.10):
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    rpr.extend([color, size])
    text = OxmlElement("w:t")
    text.text = "1"
    r.extend([rpr, text])
    fld.append(r)
    paragraph._p.append(fld)


def set_paragraph_border(paragraph, side, color=BORDER, size="6", space="4"):
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    borders.append(edge)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    for tag, value in (("tblW", USABLE_DXA), ("tblInd", TABLE_INDENT_DXA)):
        node = tblpr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tblpr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:color"), BORDER)

    old_grid = table._tbl.tblGrid
    table._tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    table._tbl.insert(1, grid)

    for row_index, row in enumerate(table.rows):
        row.height = None
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        if row_index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            trpr.append(repeat)
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tcpr = cell._tc.get_or_add_tcPr()
            tcwidth = tcpr.find(qn("w:tcW"))
            if tcwidth is None:
                tcwidth = OxmlElement("w:tcW")
                tcpr.append(tcwidth)
            tcwidth.set(qn("w:w"), str(widths[col_index]))
            tcwidth.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            if row_index == 0:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), LIGHT_GRAY)
                tcpr.append(shd)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_font(run, size=8.5 if len(widths) >= 6 else 9.5,
                             bold=True if row_index == 0 else None)


def format_document(source, destination):
    doc = Document(source)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    # Preset standard_business_brief.
    set_style(doc.styles["Normal"], 11, after=6, line=1.10)
    for name in ("Body Text", "First Paragraph"):
        if name in doc.styles:
            set_style(doc.styles[name], 11, after=6, line=1.10)
    if "Block Text" in doc.styles:
        set_style(doc.styles["Block Text"], 10.5, color=DARK_BLUE, after=8, line=1.10)
    if "Compact" in doc.styles:
        set_style(doc.styles["Compact"], 11, after=4, line=1.10)
    set_style(doc.styles["Heading 1"], 16, color=BLUE, bold=True, before=16, after=8)
    set_style(doc.styles["Heading 2"], 13, color=BLUE, bold=True, before=12, after=6)
    set_style(doc.styles["Heading 3"], 12, color=DARK_BLUE, bold=True, before=8, after=4)
    if "Image Caption" in doc.styles:
        set_style(doc.styles["Image Caption"], 9, color=MUTED, after=8, line=1.0)

    # Portada editorial sobria.
    title, subtitle = doc.paragraphs[0], doc.paragraphs[1]
    title.alignment = subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.space_after = Pt(28)
    for run in title.runs:
        set_font(run, size=28, color=NAVY, bold=True)
    for run in subtitle.runs:
        set_font(run, size=14, color=BLUE, bold=True)
    institution = subtitle.add_run("\nPontificia Universidad Católica de Chile")
    set_font(institution, size=10.5, color=MUTED, bold=False)
    for paragraph in doc.paragraphs[2:7]:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        for run in paragraph.runs:
            set_font(run, size=11, color=MUTED, bold=paragraph is doc.paragraphs[6])
    doc.paragraphs[6].paragraph_format.space_before = Pt(20)
    doc.paragraphs[6].runs[-1].add_break(WD_BREAK.PAGE)

    # Encabezado y pie, distintos en portada.
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "SERIES DE TIEMPO  |  TRABAJO FINAL"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    for run in hp.runs:
        set_font(run, size=8.5, color=MUTED, bold=True)
    set_paragraph_border(hp, "bottom", color=BORDER, size="4", space="3")
    add_page_field(section.footer.paragraphs[0])
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    # Ritmo, encabezados, imágenes, leyendas y advertencia.
    for paragraph in doc.paragraphs[7:]:
        paragraph.paragraph_format.widow_control = True
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if paragraph.style.name == "Captioned Figure":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(2)
        elif paragraph.style.name == "Image Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            for run in paragraph.runs:
                set_font(run, size=9, color=MUTED, italic=True)
        elif paragraph.style.name == "Block Text":
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.right_indent = Inches(0.08)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(10)
            set_paragraph_border(paragraph, "left", color=BLUE, size="14", space="8")
            shade_paragraph(paragraph, CALLOUT)

    # Limitar imágenes al ancho útil y mantener proporción.
    for shape in doc.inline_shapes:
        max_width = Inches(6.3)
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    width_patterns = {
        2: [2700, 6660],
        4: [2400, 2320, 2320, 2320],
        6: [2500, 1372, 1372, 1372, 1372, 1372],
    }
    for table in doc.tables:
        set_table_geometry(table, width_patterns[len(table.columns)])

    doc.core_properties.title = "Pronóstico semanal de concentración de NOx en Angamos 1"
    doc.core_properties.subject = "Trabajo final de Series de Tiempo"
    doc.core_properties.author = (
        "Hans Engelmann; Jessica Anaid Aguilar Mejía; "
        "Matías Nicolás García Garcete; Guillermo Eder López Rojas"
    )
    doc.core_properties.keywords = "series de tiempo, ARIMA, NOx, Angamos, SNIFA"
    doc.save(destination)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Uso: preparar_word_final.py entrada.docx salida.docx")
    format_document(Path(sys.argv[1]), Path(sys.argv[2]))
